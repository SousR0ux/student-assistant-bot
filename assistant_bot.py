# -*- coding: utf-8 -*-

import os
import io
import csv
import time
import html
import random
import logging
import asyncio
from datetime import datetime, timedelta
from typing import Optional, List, Dict, Any, Tuple

# Third-party libraries
import httpx
from docx import Document
from docx.shared import RGBColor
from telegram import Update, InlineKeyboardMarkup, InlineKeyboardButton, InputFile
from telegram.ext import (
    Application, CommandHandler, ContextTypes, ConversationHandler,
    MessageHandler, CallbackQueryHandler, filters, PicklePersistence
)
from telegram.helpers import escape_markdown

# ===== Load .env / .evn =====
try:
    from dotenv import load_dotenv, find_dotenv
    load_dotenv(find_dotenv(filename=".env", raise_error_if_not_found=False))
    if os.path.exists(".evn"):
        load_dotenv(".evn")
except ImportError:
    print("dotenv library not found, skipping .env file loading.")
    pass

# ===== CONFIGURATION =====
BOT_TOKEN = os.getenv("BOT_TOKEN")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
ADMIN_CHAT_ID = os.getenv("ADMIN_CHAT_ID")
FREE_LIMIT = int(os.getenv("FREE_LIMIT", "5"))
FILE_REWRITE_LIMIT = int(os.getenv("FILE_REWRITE_LIMIT", "0"))
RL_WINDOW_SEC = int(os.getenv("RL_WINDOW_SEC", "10"))
RL_MAX_HITS = int(os.getenv("RL_MAX_HITS", "3"))
CAPTCHA_ENABLED = os.getenv("CAPTCHA_ENABLED", "1") == "1"

# Referral System Config
REF_BONUS_DAYS = int(os.getenv("REF_BONUS_DAYS", "1"))
REF_WELCOME_ATTEMPTS = int(os.getenv("REF_WELCOME_ATTEMPTS", "2"))

# AI Model Config
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-1.5-flash")
GEMINI_FALLBACK_MODELS = [
    m.strip() for m in os.getenv(
        "GEMINI_FALLBACK_MODELS",
        "gemini-1.5-pro, gemini-1.0-pro"
    ).split(",")
    if m.strip() and m.strip() != GEMINI_MODEL
]
GEMINI_MAX_RETRIES = int(os.getenv("GEMINI_MAX_RETRIES", "3"))
GEMINI_BACKOFF_BASE = float(os.getenv("GEMINI_BACKOFF_BASE", "1.5"))

def _parse_admin(s: Optional[str]) -> Optional[int]:
    try:
        return int(s) if s else None
    except (ValueError, TypeError):
        return None

ADMIN_USER_ID: Optional[int] = _parse_admin(ADMIN_CHAT_ID)

# ===== LOGGING =====
logging.basicConfig(format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
                    level=logging.INFO)
logger = logging.getLogger(__name__)

# ===== CONVERSATION STATES =====
(
    MAIN_MENU, REWRITER_TEXT_INPUT, LITERATURE_TOPIC_INPUT, GOST_MENU,
    ADMIN_MENU, ADMIN_RESET_WAIT_ID, ADMIN_ADDSUB_WAIT_ID, ADMIN_ADDSUB_WAIT_DAYS,
    ADMIN_STATUS_WAIT_ID, ADMIN_DELSUB_WAIT_ID,
    CABINET_MENU, CAPTCHA_WAIT,
    SETTINGS_MENU, SETTINGS_TONE_WAIT, SETTINGS_GOST_WAIT,
    ADMIN_SEARCH_WAIT_QUERY, ADMIN_TAGS_WAIT_ID, ADMIN_TAGS_WAIT_VALUE,
    ADMIN_BROADCAST_SEGMENT, ADMIN_BROADCAST_WAIT_TEXT,
    ADMIN_SETLIMIT_WAIT_ID, ADMIN_SETLIMIT_WAIT_VALUES,
    ADMIN_BLACKLIST_WAIT_ID, ADMIN_SHADOW_WAIT_ID,
    ADMIN_METRICS_MENU,
    FILE_REWRITE_WAIT_FILE,
    ADMIN_ADDSUB_FILE_WAIT_ID, ADMIN_ADDSUB_FILE_WAIT_DAYS, ADMIN_MAINT_MSG_WAIT
) = range(29)

# ===== HELPER & UTILITY FUNCTIONS =====

# --- Date Helpers ---
def _today() -> str: return datetime.now().strftime("%Y-%m-%d")
def _now_hms() -> str: return datetime.now().strftime("%H:%M:%S")
def _ymd(dt: datetime) -> str: return dt.strftime("%Y-%m-%d")

# --- User & Permission Helpers ---
def is_admin(uid: int) -> bool:
    return ADMIN_USER_ID is not None and uid == ADMIN_USER_ID

def has_active_subscription(context: ContextTypes.DEFAULT_TYPE) -> bool:
    exp = context.user_data.get("subscription_expires")
    if not exp: return False
    try:
        return datetime.strptime(exp, "%Y-%m-%d").date() >= datetime.now().date()
    except Exception:
        return False

def has_file_rewrite_access(context: ContextTypes.DEFAULT_TYPE) -> bool:
    exp = context.user_data.get("file_subscription_expires")
    if not exp: return False
    try:
        return datetime.strptime(exp, "%Y-%m-%d").date() >= datetime.now().date()
    except Exception:
        return False

# --- Usage & Limit Helpers ---
def get_user_usage(feature: str, context: ContextTypes.DEFAULT_TYPE) -> int:
    u = context.user_data.setdefault("usage", {})
    d = u.get(feature, {"count": 0, "date": _today()})
    if d.get("date") != _today():
        d = {"count": 0, "date": _today()}
        u[feature] = d
    return int(d.get("count", 0))

def increment_usage(feature: str, context: ContextTypes.DEFAULT_TYPE) -> int:
    u = context.user_data.setdefault("usage", {})
    d = u.get(feature, {"count": 0, "date": _today()})
    if d.get("date") != _today():
        d = {"count": 1, "date": _today()}
    else:
        d["count"] = int(d.get("count", 0)) + 1
    u[feature] = d
    
    # Global stats
    stat = context.application.bot_data.setdefault("feature_usage_today", {})
    day_map = stat.setdefault(_today(), {"rewriter": 0, "literature": 0, "file_rewrite": 0})
    day_map[feature] = day_map.get(feature, 0) + 1
    return d["count"]

def remaining_attempts(feature: str, context: ContextTypes.DEFAULT_TYPE, uid: int) -> str:
    if is_admin(uid):
        return "∞ (Админ)"
    
    if feature == "file_rewrite":
        if has_file_rewrite_access(context):
            return "∞ (Подписка)"
        limit = FILE_REWRITE_LIMIT
    else:
        if has_active_subscription(context):
            return "∞ (Подписка)"
        limit = FREE_LIMIT
        
    return str(max(0, limit - get_user_usage(feature, context)))

# --- History & Analytics Helpers ---
def _push_history(context: ContextTypes.DEFAULT_TYPE, feature: str, size: int) -> None:
    hist: List[Dict[str, Any]] = context.user_data.setdefault("history", [])
    hist.append({"ts": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                 "feature": feature, "size": int(size)})
    if len(hist) > 100: del hist[:-100]

def _format_history(context: ContextTypes.DEFAULT_TYPE, limit: int = 10) -> str:
    hist: List[Dict[str, Any]] = context.user_data.get("history", [])
    if not hist: return "Пока нет записей об использовании."
    return "\n".join(f"• {i['ts']}: {i['feature']} (длина ввода: {i['size']})"
                       for i in list(reversed(hist))[:limit])

def _record_ai_stat(application: Application, ok: bool) -> None:
    stats = application.bot_data.setdefault("ai_stats", [])
    stats.append({"ts": int(time.time()), "ok": bool(ok)})
    if len(stats) > 200: del stats[:-200]

def _service_snapshot(application: Application, n: int = 50) -> Dict[str, Any]:
    stats = application.bot_data.get("ai_stats", [])
    subset = stats[-n:] if n > 0 else stats[:]
    if not subset: return {"count": 0, "error_rate": 0.0}
    errs = sum(1 for x in subset if not x.get("ok"))
    return {"count": len(subset), "error_rate": round(errs / len(subset) * 100, 1)}

def _track_active(app: Application, uid: int) -> None:
    bd = app.bot_data
    day = _today()
    dau = bd.setdefault("dau", {})
    day_set = dau.setdefault(day, set())
    day_set.add(uid)
    allu = bd.setdefault("all_users", set())
    allu.add(uid)

def _count_dau(bd: dict, days: int) -> int:
    today = datetime.now().date()
    dau: Dict[str, set] = bd.get("dau", {})
    seen: set = set()
    for i in range(days):
        d_str = (today - timedelta(days=i)).strftime("%Y-%m-%d")
        if d_str in dau:
            seen.update(dau[d_str])
    return len(seen)

def _count_new_users(app: Application, days: int) -> int:
    cutoff = datetime.now().date() - timedelta(days=days - 1)
    c = 0
    for _, ud in app.user_data.items():
        fs = ud.get("first_seen")
        if fs:
            try:
                d = datetime.strptime(fs, "%Y-%m-%d").date()
                if d >= cutoff:
                    c += 1
            except (ValueError, TypeError):
                continue
    return c

def _feature_usage_today(bd: dict) -> Dict[str, int]:
    fm = bd.get("feature_usage_today", {}).get(_today(), {})
    return {
        "rewriter": int(fm.get("rewriter", 0)),
        "literature": int(fm.get("literature", 0)),
        "file_rewrite": int(fm.get("file_rewrite", 0))
    }

# --- Security & Anti-Spam Helpers ---
def _ensure_first_seen(context: ContextTypes.DEFAULT_TYPE) -> None:
    if "first_seen" not in context.user_data:
        context.user_data["first_seen"] = _today()

def _touch_seen(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    context.user_data["last_seen"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    u = update.effective_user
    if u:
        context.user_data["last_username"] = (u.username or "")[:64]
        context.user_data["first_name"] = (u.first_name or "")[:128]
        context.user_data["last_name"] = (u.last_name or "")[:128]
    _track_active(context.application, update.effective_user.id)
    _ensure_first_seen(context)

def _is_blacklisted(app: Application, uid: int) -> bool:
    bl = app.bot_data.get("blacklist", set())
    return uid in bl

def _is_shadowbanned(app: Application, uid: int) -> bool:
    sb = app.bot_data.get("shadowban", set())
    return uid in sb

def _rate_limit_ok(context: ContextTypes.DEFAULT_TYPE) -> Tuple[bool, int]:
    now = time.time()
    arr: List[float] = context.user_data.setdefault("rl_times", [])
    arr = [t for t in arr if now - t <= RL_WINDOW_SEC]
    ok = len(arr) < RL_MAX_HITS
    if ok:
        arr.append(now)
        context.user_data["rl_times"] = arr
        return True, 0
    wait = RL_WINDOW_SEC - int(now - min(arr)) if arr else RL_WINDOW_SEC
    return False, max(wait, 1)


def _need_captcha(context: ContextTypes.DEFAULT_TYPE, uid: int) -> bool:
    if not CAPTCHA_ENABLED or is_admin(uid):
        return False
    return not context.user_data.get("captcha_ok", False)

def _gen_captcha(context: ContextTypes.DEFAULT_TYPE) -> str:
    a, b = random.randint(2, 9), random.randint(2, 9)
    context.user_data["captcha_answer"] = str(a + b)
    return f"Проверка: сколько будет {a} + {b}? Отправьте ответ числом."

# --- Maintenance Mode Helpers ---
MAINTENANCE_DEFAULT_MSG = (
    "🛠 <b>Сервис временно на техработах</b>.\n"
    "Попробуйте позже. Если вопрос срочный — напишите автору: "
    "<a href='https://t.me/V_L_A_D_IS_L_A_V'>@V_L_A_D_IS_L_A_V</a>"
)

def _maintenance_on(app: Application) -> bool:
    return bool(app.bot_data.get("maintenance_enabled", False))

def _maintenance_text(app: Application) -> str:
    txt = app.bot_data.get("maintenance_msg")
    return txt if (isinstance(txt, str) and txt.strip()) else MAINTENANCE_DEFAULT_MSG

async def _maintenance_guard(update: Update, context: ContextTypes.DEFAULT_TYPE) -> bool:
    if _maintenance_on(context.application) and not is_admin(update.effective_user.id):
        msg = _maintenance_text(context.application)
        if update.callback_query:
            await update.callback_query.answer()
            await update.callback_query.message.edit_text(
                msg, parse_mode="HTML", reply_markup=contact_kb(), disable_web_page_preview=True
            )
        else:
            await update.effective_message.reply_html(
                msg, reply_markup=contact_kb(), disable_web_page_preview=True
            )
        return True
    return False

# --- Text Formatting Helpers ---
TG_MD_LIMIT = 3800
def _chunk_md(text: str, limit: int = TG_MD_LIMIT) -> List[str]:
    chunks, buf = [], ""
    for line in text.splitlines(keepends=True):
        if len(buf) + len(line) > limit:
            if buf: chunks.append(buf)
            buf = line
        else:
            buf += line
    if buf: chunks.append(buf)
    return chunks or ["_пусто_"]

async def _md_send_chunks(msg, text: str, markup=None):
    parts = _chunk_md(text, TG_MD_LIMIT)
    if not parts: return

    await msg.edit_text(
        parts[0],
        parse_mode="Markdown",
        disable_web_page_preview=True,
        reply_markup=markup if len(parts) == 1 else None
    )

    if len(parts) > 1:
        for part in parts[1:-1]:
            await msg.reply_text(
                part,
                parse_mode="Markdown",
                disable_web_page_preview=True
            )
        await msg.reply_text(
            parts[-1],
            parse_mode="Markdown",
            disable_web_page_preview=True,
            reply_markup=markup
        )


def _progress_bar(used: int, total: int, width: int = 20) -> str:
    if total <= 0: return "░" * width + " ∞"
    used = max(0, min(total, used))
    fill = int(width * used / total)
    return "█" * fill + "░" * (width - fill) + f" {used}/{total}"

def _no_literature_found(txt: str) -> bool:
    s = (txt or "").strip().lower()
    if not s or s.startswith(("ошибка", "не удалось")): return True
    if any(x in s for x in ["ничего не найдено", "источники не найд", "нет данных", "no sources"]): return True
    if s.count("\n") <= 1 and len(s) < 60: return True
    return False

# ===== AI SERVICE (GEMINI) - IMPROVED ERROR HANDLING =====
async def call_gemini(prompt: str) -> str:
    if not GEMINI_API_KEY:
        return "Ошибка: API-ключ для нейросети не настроен."

    async def _try_call(model: str) -> tuple[bool, str]:
        api_url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={GEMINI_API_KEY}"
        payload = {"contents": [{"parts": [{"text": prompt}]}]}
        try:
            async with httpx.AsyncClient(timeout=90.0) as client:
                r = await client.post(api_url, json=payload)

                if r.status_code != 200:
                    error_info = r.json()
                    error_message = error_info.get("error", {}).get("message", r.text)
                    logger.error(f"Gemini API error (HTTP {r.status_code}) on model {model}: {error_message}")
                    if r.status_code == 429: return False, "quota_or_rate"
                    if 500 <= r.status_code < 600: return False, f"server_{r.status_code}"
                    # For other client errors (4xx), it's likely a persistent issue.
                    return True, f"Ошибка API: {error_message}"

                j = r.json()
                text = j.get("candidates", [{}])[0].get("content", {}).get("parts", [{}])[0].get("text", "").strip()
                return (True, text) if text else (False, "empty_answer")
        except (httpx.ConnectError, httpx.ReadTimeout):
            logger.warning(f"Gemini call network error on model {model}")
            return False, "network_error"
        except Exception as e:
            logger.error(f"Gemini call unexpected error on model {model}", exc_info=True)
            return True, f"Произошла внутренняя ошибка при обращении к нейросети."

    models_chain = [GEMINI_MODEL] + GEMINI_FALLBACK_MODELS
    last_error_message = "Не удалось получить ответ от нейросети."

    for model in models_chain:
        for attempt in range(GEMINI_MAX_RETRIES):
            ok, response_text = await _try_call(model)
            if ok:
                return response_text
            
            last_error_message = response_text # Store the reason for retry
            if response_text in ("quota_or_rate", "network_error") or response_text.startswith("server_"):
                delay = (GEMINI_BACKOFF_BASE ** attempt) + random.uniform(0, 0.5)
                await asyncio.sleep(delay)
                continue
            else:
                break 
                
    logger.error(f"AI call failed after all retries and fallbacks. Last error: {last_error_message}")
    return "Сервис ИИ временно недоступен. Пожалуйста, попробуйте позже."


# ===== FILE REWRITER UTILS =====
async def read_telegram_file(update: Update) -> tuple[bytes, str]:
    file = await update.message.document.get_file()
    file_bytes = await file.download_as_bytearray()
    return bytes(file_bytes), update.message.document.file_name

def process_docx_for_rewrite(file_bytes: bytes) -> tuple[Document, List[Dict[str, Any]]]:
    doc = Document(io.BytesIO(file_bytes))
    highlighted_parts = []
    
    for para_idx, paragraph in enumerate(doc.paragraphs):
        for run_idx, run in enumerate(paragraph.runs):
            if run.font.highlight_color is not None:
                highlighted_parts.append({
                    "text": run.text,
                    "para_idx": para_idx,
                    "run_idx": run_idx,
                })
    return doc, highlighted_parts

async def rewrite_highlighted_parts_async(parts: List[Dict[str, Any]], rewrite_fn, tone: str, **kwargs) -> List[Dict[str, Any]]:
    tasks = []
    for part in parts:
        prompt = (
            f"Перепиши следующий фрагмент текста в {tone} стиле, сохраняя основной смысл. "
            "Не добавляй вступлений, выводов или лишних фраз. "
            f"Исходный текст:\n\n\"{part['text']}\""
        )
        tasks.append(rewrite_fn(prompt))
    
    rewritten_texts = await asyncio.gather(*tasks)
    
    for i, part in enumerate(parts):
        part["rewritten_text"] = rewritten_texts[i]
        
    return parts

def build_final_docx(original_doc: Document, rewritten_parts: List[Dict[str, Any]]) -> bytes:
    for part in rewritten_parts:
        para_idx, run_idx = part["para_idx"], part["run_idx"]
        
        if para_idx < len(original_doc.paragraphs):
            paragraph = original_doc.paragraphs[para_idx]
            if run_idx < len(paragraph.runs):
                run = paragraph.runs[run_idx]
                run.text = part["rewritten_text"]
                run.font.highlight_color = None 
                run.font.color.rgb = RGBColor(0, 0, 0)

    bio = io.BytesIO()
    original_doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


# ===== KEYBOARDS =====
def main_menu_kb(uid: int) -> InlineKeyboardMarkup:
    rows = [
        [InlineKeyboardButton("✍️ AI-Рерайтер текста", callback_data="rewriter")],
        [InlineKeyboardButton("📄 AI-Рерайт файла (DOCX)", callback_data="file_rewriter")],
        [InlineKeyboardButton("📚 Генератор списка литературы", callback_data="literature")],
        [InlineKeyboardButton("📋 Консультант по ГОСТу", callback_data="gost")],
        [InlineKeyboardButton("👤 Личный кабинет", callback_data="cabinet")],
    ]
    if is_admin(uid):
        rows.append([InlineKeyboardButton("⚙️ Админ-панель", callback_data="admin_panel")])
    return InlineKeyboardMarkup(rows)

def back_menu_kb() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup([[InlineKeyboardButton("⬅️ Назад в меню", callback_data="back_to_main_menu")]])

def contact_kb() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("💬 Написать автору", url="https://t.me/V_L_A_D_IS_L_A_V")],
        [InlineKeyboardButton("⬅️ Назад в меню", callback_data="back_to_main_menu")]
    ])

def cabinet_kb() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("🗂 История (10)", callback_data="cab_history")],
        [InlineKeyboardButton("📥 Экспорт истории (CSV)", callback_data="cab_export")],
        [InlineKeyboardButton("⚙️ Настройки", callback_data="cab_settings")],
        [InlineKeyboardButton("👥 Рефералы", callback_data="cab_ref")],
        [InlineKeyboardButton("🔥 Получить безлимит", url="https://t.me/V_L_A_D_IS_L_A_V")],
        [InlineKeyboardButton("🔄 Обновить", callback_data="cab_refresh")],
        [InlineKeyboardButton("⬅️ Назад в меню", callback_data="back_to_main_menu")],
    ])

def admin_menu_kb() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("🔁 Сбросить лимиты", callback_data="admin_reset"),
         InlineKeyboardButton("📊 Статус", callback_data="admin_status")],
        [InlineKeyboardButton("➕ Подписка (общая)", callback_data="admin_addsub"),
         InlineKeyboardButton("📄 Подписка (файлы)", callback_data="admin_addsub_file")],
        [InlineKeyboardButton("❌ Отменить подписку", callback_data="admin_delsub")],
        [InlineKeyboardButton("🔎 Поиск", callback_data="admin_search"),
         InlineKeyboardButton("🏷 Теги", callback_data="admin_tags")],
        [InlineKeyboardButton("📣 Рассылка", callback_data="admin_broadcast"),
         InlineKeyboardButton("📤 Экспорт CSV", callback_data="admin_export")],
        [InlineKeyboardButton("🚫 Блокировка", callback_data="admin_blacklist"),
         InlineKeyboardButton("👻 Теневой бан", callback_data="admin_shadow")],
        [InlineKeyboardButton("🎚 Задать лимиты", callback_data="admin_setlimit"),
         InlineKeyboardButton("📈 Метрики", callback_data="admin_metrics")],
         [InlineKeyboardButton("🛠 Тех. работы", callback_data="admin_maint_toggle"),
          InlineKeyboardButton("📝 Текст тех. работ", callback_data="admin_maint_msg")],
        [InlineKeyboardButton("⬅️ Назад в главное меню", callback_data="back_to_main_menu")],
    ])

def admin_cancel_kb() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("⬅️ Назад в админ-панель", callback_data="admin_panel")],
        [InlineKeyboardButton("🏠 В меню", callback_data="back_to_main_menu")],
    ])

def settings_kb() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("🗣 Тон рерайта", callback_data="set_tone")],
        [InlineKeyboardButton("📏 ГОСТ по умолчанию", callback_data="set_gost")],
        [InlineKeyboardButton("⬅️ Назад", callback_data="cabinet")],
    ])

def admin_metrics_kb() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("🔄 Обновить", callback_data="admin_metrics_refresh"),
         InlineKeyboardButton("📥 Экспорт (CSV)", callback_data="admin_metrics_export")],
        [InlineKeyboardButton("⬅️ Назад", callback_data="admin_panel")],
    ])

def broadcast_segments_kb() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("Все", callback_data="bseg_all"),
         InlineKeyboardButton("Free", callback_data="bseg_free"),
         InlineKeyboardButton("Pro", callback_data="bseg_subs")],
        [InlineKeyboardButton("Неактивные 7д", callback_data="bseg_inactive")],
        [InlineKeyboardButton("🔇 Все (тихо)", callback_data="bseg_all_silent")],
        [InlineKeyboardButton("🔇 Free (тихо)", callback_data="bseg_free_silent")],
        [InlineKeyboardButton("🔇 Pro (тихо)", callback_data="bseg_subs_silent")],
        [InlineKeyboardButton("⬅️ Назад", callback_data="admin_panel")]
    ])

# ===== BOT HANDLERS (omitted for brevity, they are the same as the previous correct version) =====
# ... All handlers from start() to error_handler() are placed here ...
# For the final file, the full set of handlers from the previous correct version is included.
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    try:
        if context.args:
            ref_from = int(context.args[0])
            me = update.effective_user.id
            if me != ref_from and not context.user_data.get("is_ref"):
                ref_ud = context.application.user_data.get(ref_from)
                if isinstance(ref_ud, dict):
                    refs = ref_ud.setdefault("referrals", set())
                    if me not in refs:
                        refs.add(me)
                        context.user_data["is_ref"] = True
                        if REF_BONUS_DAYS > 0:
                            start_date = datetime.now().date()
                            exp_str = ref_ud.get("subscription_expires")
                            if exp_str:
                                try:
                                    exp_date = datetime.strptime(exp_str, "%Y-%m-%d").date()
                                    if exp_date > start_date: start_date = exp_date
                                except (ValueError, TypeError): pass
                            until = start_date + timedelta(days=REF_BONUS_DAYS)
                            ref_ud["subscription_expires"] = until.strftime("%Y-%m-%d")
                            try:
                                await context.bot.send_message(
                                    chat_id=ref_from,
                                    text=f"🎁 Новый реферал! Вам продлён безлимит до {until.strftime('%d.%m.%Y')}."
                                )
                            except Exception: pass
                        if REF_WELCOME_ATTEMPTS > 0:
                            usage = context.user_data.setdefault("usage", {})
                            for feat in ("rewriter", "literature"):
                                d = usage.get(feat, {"count": 0, "date": _today()})
                                d["count"] = max(0, d.get("count", 0) - REF_WELCOME_ATTEMPTS)
                                usage[feat] = d
    except Exception:
        pass

    _touch_seen(update, context)
    uid = update.effective_user.id

    if _is_blacklisted(context.application, uid):
        await update.effective_message.reply_html("🚫 Доступ к боту ограничен.")
        return ConversationHandler.END
    if _is_shadowbanned(context.application, uid):
        await update.effective_message.reply_text("Сервис временно перегружен, попробуйте позже.")
        return ConversationHandler.END

    if _need_captcha(context, uid):
        q = _gen_captcha(context)
        await update.effective_message.reply_text(q)
        return CAPTCHA_WAIT

    text = f"Здравствуйте, {update.effective_user.mention_html()}!\n\nЯ «Студент-Ассистент». Выберите инструмент:"
    if update.callback_query:
        await update.callback_query.answer()
        await update.callback_query.message.edit_text(text, parse_mode="HTML", reply_markup=main_menu_kb(uid))
    else:
        await update.effective_message.reply_html(text, reply_markup=main_menu_kb(uid))
    return MAIN_MENU

async def captcha_check(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    ans = (update.message.text or "").strip()
    if ans == context.user_data.get("captcha_answer"):
        context.user_data["captcha_ok"] = True
        await update.message.reply_text("✅ Спасибо! Продолжаем.")
        return await start(update, context)
    await update.message.reply_text("❌ Неверно. " + _gen_captcha(context))
    return CAPTCHA_WAIT

async def help_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    t = (
        "📖 Помощь\n\n"
        "• /start — главное меню\n"
        "• /status — ваш доступ и остаток попыток\n"
        "• /report <сообщение> — жалоба с логами\n"
        "• /service — статус ИИ\n\n"
        "Для безлимита и вопросов: @V_L_A_D_IS_L_A_V"
    )
    await update.message.reply_text(t)

async def service_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    snap = _service_snapshot(context.application, 50)
    await update.message.reply_html(
        f"🩺 <b>Статус</b>\nПоследних запросов: {snap['count']}\nОшибка ИИ: {snap['error_rate']}%\n"
        "Если ошибка высока — возможны задержки/сбои."
    )

async def report_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    msg = (update.message.text or "").partition(" ")[2].strip()
    if not ADMIN_USER_ID:
        await update.message.reply_text("Администратор не настроен.")
        return
    
    last_req = context.user_data.get("last_request", {})
    hist = _format_history(context, limit=5)
    uname = f"@{(update.effective_user.username or '').strip()}" if update.effective_user.username else "—"
    html_msg = (
        f"🛎 <b>Жалоба</b>\n"
        f"ID: <code>{update.effective_user.id}</code>\n"
        f"Юзернейм: {html.escape(uname)}\n"
        f"Текст: {html.escape(msg) or '—'}\n\n"
        f"<b>Last request:</b> {html.escape(str(last_req))}\n\n"
        f"<b>History(5):</b>\n{html.escape(hist)}"
    )
    try:
        await context.bot.send_message(chat_id=ADMIN_USER_ID, text=html_msg, parse_mode="HTML")
        await update.message.reply_text("Спасибо! Сообщение отправлено.")
    except Exception as e:
        logger.warning("report send fail: %s", e)
        await update.message.reply_text("Не удалось отправить отчёт. Напишите @V_L_A_D_IS_L_A_V")

async def check_status(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    uid = update.effective_user.id
    if has_active_subscription(context):
        exp = datetime.strptime(context.user_data.get("subscription_expires"), "%Y-%m-%d").strftime("%d.%m.%Y")
        await update.message.reply_html(f"<b>Статус:</b> ✅ Безлимит до {exp}")
    else:
        r = remaining_attempts("rewriter", context, uid)
        l = remaining_attempts("literature", context, uid)
        await update.message.reply_html(
            f"<b>Статус:</b> базовый доступ\n"
            f"• Рерайтер: {r} из {FREE_LIMIT}\n"
            f"• Литература: {l} из {FREE_LIMIT}"
        )

async def rewriter_start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    q = update.callback_query
    uid = update.effective_user.id
    await q.answer()

    if await _maintenance_guard(update, context): return MAIN_MENU
    
    if not is_admin(uid) and not has_active_subscription(context):
        if get_user_usage("rewriter", context) >= FREE_LIMIT:
            await q.message.edit_text(
                "🚫 <b>Дневной лимит исчерпан</b>\n\nДля безлимита напишите: <a href='https://t.me/V_L_A_D_IS_L_A_V'>@V_L_A_D_IS_L_A_V</a>",
                parse_mode="HTML", reply_markup=contact_kb()
            )
            return MAIN_MENU

    left = remaining_attempts("rewriter", context, uid)
    await q.message.edit_text(
        f"✍️ *AI-Рерайтер*\n\nПришлите текст (до 2000 симв.).\n\nДоступно сегодня: *{left}*",
        parse_mode="Markdown", reply_markup=back_menu_kb()
    )
    return REWRITER_TEXT_INPUT

async def rewriter_process_text(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    if await _maintenance_guard(update, context): return REWRITER_TEXT_INPUT

    _touch_seen(update, context)
    uid = update.effective_user.id
    ok_rl, wait = _rate_limit_ok(context)
    if not ok_rl:
        await update.message.reply_text(f"⏳ Слишком часто. Попробуйте через {wait} сек.")
        return REWRITER_TEXT_INPUT

    if not is_admin(uid) and not has_active_subscription(context):
        if get_user_usage("rewriter", context) >= FREE_LIMIT:
            await update.message.reply_html(
                "🚫 <b>Дневной лимит исчерпан.</b>\n\nДля безлимита напишите: <a href='https://t.me/V_L_A_D_IS_L_A_V'>@V_L_A_D_IS_L_A_V</a>",
                reply_markup=contact_kb()
            )
            return REWRITER_TEXT_INPUT

    user_text = (update.message.text or "")[:2000]
    context.user_data["last_request"] = {"feature": "rewriter", "len": len(user_text), "ts": datetime.now().isoformat()}

    processing = await update.message.reply_text("⏳ Обрабатываю…")
    tone = context.user_data.get("tone", "официальный")
    prompt = (
        f"Перепиши текст в {tone} стиле, сохранив смысл и структуру. Не добавляй вступлений/выводов. "
        f"Исходный текст:\n\"\"\"\n{user_text}\n\"\"\""
    )
    txt = await call_gemini(prompt)

    success = not txt.startswith("Ошибка") and "Сервис ИИ временно недоступен" not in txt
    _record_ai_stat(context.application, success)
    if success:
        if not is_admin(uid) and not has_active_subscription(context):
            increment_usage("rewriter", context)
        _push_history(context, "rewriter", len(user_text))

    left = remaining_attempts("rewriter", context, uid)
    full = f"*Готово! Ваш текст:*\n\n{txt}\n\n*Доступно сегодня:* {left}"
    await _md_send_chunks(processing, full, markup=back_menu_kb())
    return REWRITER_TEXT_INPUT

async def literature_start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    q = update.callback_query
    uid = update.effective_user.id
    await q.answer()

    if await _maintenance_guard(update, context): return MAIN_MENU

    if not is_admin(uid) and not has_active_subscription(context) and get_user_usage("literature", context) >= FREE_LIMIT:
        await q.message.edit_text(
            "🚫 <b>Дневной лимит исчерпан.</b>\n\nДля безлимита напишите: <a href='https://t.me/V_L_A_D_IS_L_A_V'>@V_L_A_D_IS_L_A_V</a>",
            parse_mode="HTML", reply_markup=contact_kb()
        )
        return MAIN_MENU

    left = remaining_attempts("literature", context, uid)
    await q.message.edit_text(
        f"📚 *Генератор списка литературы*\n\nНапишите тему.\n\nДоступно сегодня: *{left}*",
        parse_mode="Markdown", reply_markup=back_menu_kb()
    )
    return LITERATURE_TOPIC_INPUT

async def literature_process_topic(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    if await _maintenance_guard(update, context): return LITERATURE_TOPIC_INPUT

    _touch_seen(update, context)
    uid = update.effective_user.id
    ok_rl, wait = _rate_limit_ok(context)
    if not ok_rl:
        await update.message.reply_text(f"⏳ Слишком часто. Попробуйте через {wait} сек.")
        return LITERATURE_TOPIC_INPUT

    if not is_admin(uid) and not has_active_subscription(context) and get_user_usage("literature", context) >= FREE_LIMIT:
        await update.message.reply_html(
            "🚫 <b>Дневной лимит исчерпан.</b>\n\nДля безлимита напишите: <a href='https://t.me/V_L_A_D_IS_L_A_V'>@V_L_A_D_IS_L_A_V</a>",
            reply_markup=contact_kb()
        )
        return LITERATURE_TOPIC_INPUT

    topic = (update.message.text or "")[:500]
    context.user_data["last_request"] = {"feature": "literature", "len": len(topic), "ts": datetime.now().isoformat()}
    processing = await update.message.reply_text("📚 Подбираю источники…")

    prompt = (
        "Ты — AI-библиограф. Составь НУМЕРОВАННЫЙ список 5–7 актуальных источников (книги, статьи) на русском языке. "
        "Для каждого укажи библиографию и короткий комментарий (1–2 предложения) о пользе источника. "
        "Строгий формат: 1. **Автор(ы), И. О.** (Год). *Название*. Город: Издательство. **Комментарий:** ...\n"
        f"Тема: «{topic}»\n"
        "Не добавляй ссылки и лишние прелюдии/итоги."
    )
    txt = await call_gemini(prompt)

    success = not _no_literature_found(txt)
    _record_ai_stat(context.application, success)
    
    if not success:
        await processing.edit_text(
            "😕 <b>Подходящие источники не нашлись.</b>\nПопробуйте сузить тему или напишите мне: @V_L_A_D_IS_L_A_V",
            parse_mode="HTML", reply_markup=contact_kb(), disable_web_page_preview=True
        )
        return LITERATURE_TOPIC_INPUT

    if not is_admin(uid) and not has_active_subscription(context):
        increment_usage("literature", context)
    _push_history(context, "literature", len(topic))

    left = remaining_attempts("literature", context, uid)
    full = f"*Готово! Рекомендуемый список:*\n\n{txt}\n\n*Доступно сегодня:* {left}"
    await _md_send_chunks(processing, full, markup=back_menu_kb())
    return LITERATURE_TOPIC_INPUT

async def file_rewriter_start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    # Отвечаем пользователю всплывающим окном
    await update.callback_query.answer(
        "Эта функция находится в разработке и скоро будет доступна.",
        show_alert=True
    )
    # Остаемся в главном меню, не переходя в другие состояния
    return MAIN_MENU

async def process_document_rewrite(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    _touch_seen(update, context)
    uid = update.effective_user.id

    if await _maintenance_guard(update, context): return FILE_REWRITE_WAIT_FILE

    if not is_admin(uid) and not has_file_rewrite_access(context):
        await update.message.reply_html("🚫 У вас нет подписки на рерайт файлов.", reply_markup=contact_kb())
        return FILE_REWRITE_WAIT_FILE

    try:
        file_bytes, filename = await read_telegram_file(update)
        if not filename.lower().endswith(".docx"):
            await update.message.reply_text("❌ Пожалуйста, отправьте файл в формате .docx")
            return FILE_REWRITE_WAIT_FILE

        processing_msg = await update.message.reply_text("⏳ Ищу выделенный текст в DOCX...")
        
        original_doc, highlighted_parts = process_docx_for_rewrite(file_bytes)
        
        if not highlighted_parts:
            await processing_msg.edit_text("⚠️ В файле не найден текст, выделенный жёлтым цветом.")
            return FILE_REWRITE_WAIT_FILE

        await processing_msg.edit_text(f"Найдено {len(highlighted_parts)} фрагмент(ов). Отправляю в AI...")

        rewritten_parts = await rewrite_highlighted_parts_async(
            highlighted_parts,
            rewrite_fn=call_gemini,
            tone=context.user_data.get("tone", "официальный")
        )

        await processing_msg.edit_text("✅ Собираю итоговый документ...")
        
        docx_bytes = build_final_docx(original_doc, rewritten_parts)
        _push_history(context, "file_rewrite", len(file_bytes))

        new_filename = f"rewritten_{os.path.splitext(filename)[0]}.docx"
        bio = io.BytesIO(docx_bytes)
        bio.name = new_filename
        
        await processing_msg.delete()
        await update.message.reply_document(
            InputFile(bio), 
            caption="Готово! Выделенные части переписаны."
        )

    except Exception as e:
        logger.error("Ошибка при обработке файла: %s", e, exc_info=True)
        await update.message.reply_text(f"❌ Произошла ошибка: {e}")

    return FILE_REWRITE_WAIT_FILE

async def gost_menu(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    q = update.callback_query; await q.answer()
    kb = [
        [InlineKeyboardButton("Как оформить сноску?", callback_data="gost_footnote")],
        [InlineKeyboardButton("Как оформить список литературы?", callback_data="gost_references")],
        [InlineKeyboardButton("Общие требования", callback_data="gost_general")],
        [InlineKeyboardButton("⬅️ Назад", callback_data="back_to_main_menu")],
    ]
    await q.message.edit_text(
        "📋 **Консультант по ГОСТу**\n\nВыберите вопрос:",
        reply_markup=InlineKeyboardMarkup(kb), parse_mode="Markdown"
    )
    return GOST_MENU

async def gost_show_info(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    q = update.callback_query
    await q.answer()
    info_type = q.data
    text = ""
    if info_type == "gost_footnote":
        text = (
            "**📄 Оформление сносок**\n\n"
            "**Книга:**\n¹ Иванов И. И. Название. – М.: Издательство, 2023. – С. 45.\n\n"
            "**Статья:**\n² Петров П. П. Название // Журнал. – 2022. – № 2. – С. 12-15.\n\n"
            "**Сайт:**\n³ Название [Электронный ресурс]. – URL: http://... (дата обращения: 23.08.2025)."
        )
    elif info_type == "gost_references":
        text = (
            "**📚 Оформление списка литературы**\n\n"
            "**Книга:**\nИванов, И. И. Название. – Москва : Издательство, 2023. – 250 с.\n\n"
            "**Статья:**\nПетров, П. П. Название // Журнал. – 2022. – № 2. – С. 12–15."
        )
    elif info_type == "gost_general":
        text = (
            "**⚙️ Общие требования**\n\n"
            "• **Шрифт:** Times New Roman, 14 пт\n"
            "• **Интервал:** 1,5\n"
            "• **Выравнивание:** по ширине\n"
            "• **Отступ:** 1,25 см\n"
            "• **Поля:** левое–3, правое–1, верх/низ–2 см.\n\n"
            "*Всегда сверяйтесь с методичкой вашего вуза.*"
        )
    await q.message.edit_text(
        text=text,
        reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("⬅️ Назад", callback_data="gost_back")]]),
        parse_mode="Markdown"
    )
    return GOST_MENU

async def cabinet_open(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    uid = update.effective_user.id
    tone = context.user_data.get("tone", "официальный")
    gost = context.user_data.get("gost", "универсальный")

    if has_active_subscription(context):
        exp = datetime.strptime(context.user_data["subscription_expires"], "%Y-%m-%d").strftime("%d.%m.%Y")
        sub_text = f"✅ Безлимит до {exp}"
    else:
        sub_text = "базовый доступ"
    
    if has_file_rewrite_access(context):
        file_exp = datetime.strptime(context.user_data["file_subscription_expires"], "%Y-%m-%d").strftime("%d.%m.%Y")
        file_sub_text = f"✅ Подписка до {file_exp}"
    else:
        file_sub_text = "нет подписки"

    next_reset = (datetime.now().date() + timedelta(days=1)).strftime('%d.%m.%Y')

    text = (
        f"👤 <b>Личный кабинет</b>\n\n"
        f"<b>ID:</b> <code>{uid}</code>\n"
        f"<b>Основной доступ:</b> {sub_text}\n"
        f"<b>Рерайт файлов:</b> {file_sub_text}\n"
        f"<b>Сброс лимитов:</b> {next_reset} 00:00\n\n"
        f"✍️ Рерайтер: {remaining_attempts('rewriter', context, uid)}\n"
        f"📚 Литература: {remaining_attempts('literature', context, uid)}\n"
        f"📄 Рерайт файлов (бесп.): {remaining_attempts('file_rewrite', context, uid)}\n\n"
        f"🗣 Тон: <b>{html.escape(tone)}</b>, 📏 ГОСТ: <b>{html.escape(gost)}</b>"
    )

    if update.callback_query:
        await update.callback_query.answer()
        await update.callback_query.message.edit_text(text, parse_mode="HTML", reply_markup=cabinet_kb(), disable_web_page_preview=True)
    else:
        await update.effective_message.reply_html(text, reply_markup=cabinet_kb(), disable_web_page_preview=True)
    return CABINET_MENU

async def cabinet_refresh(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    return await cabinet_open(update, context)

async def cabinet_history(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    msg = update.callback_query.message if update.callback_query else update.effective_message
    if update.callback_query: await update.callback_query.answer()
    await msg.edit_text("🗂 <b>Последние 10 использований</b>\n\n" + _format_history(context, 10),
                        parse_mode="HTML", reply_markup=cabinet_kb())
    return CABINET_MENU

async def cabinet_export(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    hist = context.user_data.get("history", [])
    buf = io.StringIO(); w = csv.writer(buf); w.writerow(["ts", "feature", "size"])
    for h in hist: w.writerow([h.get(k, "") for k in ["ts", "feature", "size"]])
    byte = io.BytesIO(buf.getvalue().encode("utf-8")); byte.name = "history.csv"
    await update.callback_query.answer()
    await update.callback_query.message.reply_document(InputFile(byte), caption="История (CSV)")
    return CABINET_MENU

async def cabinet_referrals(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await update.callback_query.answer()
    uid = update.effective_user.id
    refs = context.user_data.get("referrals", set())
    count = len(refs) if isinstance(refs, set) else 0
    
    bot_username = context.application.bot_data.get("bot_username")
    if not bot_username:
        me = await context.bot.get_me()
        bot_username = me.username
        context.application.bot_data["bot_username"] = bot_username
        
    link = f"https://t.me/{bot_username}?start={uid}"
    
    txt = (
        f"👥 <b>Рефералы</b>\n\n"
        f"Ваша ссылка:\n<code>{html.escape(link)}</code>\n\n"
        f"За каждого друга вы получаете <b>+{REF_BONUS_DAYS} дн.</b> безлимита.\n"
        f"У вас: <b>{count}</b> реферал(ов)."
    )
    kb = InlineKeyboardMarkup([
        [InlineKeyboardButton("⬅️ Назад", callback_data="cabinet")],
        [InlineKeyboardButton("Поделиться", url=f"https://t.me/share/url?url={link}&text=AI-помощник для учебы")]
    ])
    await update.callback_query.message.edit_text(txt, parse_mode="HTML", reply_markup=kb, disable_web_page_preview=True)
    return CABINET_MENU

async def settings_open(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await update.callback_query.answer()
    await update.callback_query.message.edit_text("⚙️ Настройки", reply_markup=settings_kb())
    return SETTINGS_MENU

async def settings_pick(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await update.callback_query.answer()
    data = update.callback_query.data
    if data == "set_tone":
        await update.callback_query.message.edit_text("Введите тон рерайта (напр. официальный):", reply_markup=admin_cancel_kb())
        return SETTINGS_TONE_WAIT
    if data == "set_gost":
        await update.callback_query.message.edit_text("Введите ГОСТ по умолчанию (напр. универсальный):", reply_markup=admin_cancel_kb())
        return SETTINGS_GOST_WAIT
    return SETTINGS_MENU

async def settings_tone_save(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    context.user_data["tone"] = (update.message.text or "официальный")[:50]
    await update.message.reply_text("✅ Тон сохранён.")
    return await cabinet_open(update, context)

async def settings_gost_save(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    context.user_data["gost"] = (update.message.text or "универсальный")[:50]
    await update.message.reply_text("✅ ГОСТ сохранён.")
    return await cabinet_open(update, context)

async def reset_limit(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    if not is_admin(update.effective_user.id): return
    try:
        target_id = int(context.args[0])
        ud = context.application.user_data.get(target_id)
        if ud:
            today = _today()
            usage = ud.setdefault("usage", {})
            usage["rewriter"] = {"count": 0, "date": today}
            usage["literature"] = {"count": 0, "date": today}
            usage["file_rewrite"] = {"count": 0, "date": today}
            await update.message.reply_text("✅ Лимиты сброшены.")
        else:
            await update.message.reply_text("Пользователь не найден.")
    except (IndexError, ValueError):
        await update.message.reply_text("Используйте: /reset <user_id>")

async def add_subscription(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    if not is_admin(update.effective_user.id): return
    try:
        target_id, days = int(context.args[0]), int(context.args[1])
        ud = context.application.user_data.get(target_id)
        if not ud: await update.message.reply_text("Пользователь не найден."); return
        
        start_date = datetime.now().date()
        exp_str = ud.get("subscription_expires")
        if exp_str:
            try:
                exp_date = datetime.strptime(exp_str, "%Y-%m-%d").date()
                if exp_date > start_date: start_date = exp_date
            except (ValueError, TypeError): pass

        new_exp = start_date + timedelta(days=days)
        ud["subscription_expires"] = new_exp.strftime("%Y-%m-%d")
        await update.message.reply_text(f"✅ Подписка для {target_id} до {new_exp.strftime('%d.%m.%Y')}")
    except (IndexError, ValueError):
        await update.message.reply_text("Используйте: /addsub <user_id> <days>")

async def del_subscription(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    if not is_admin(update.effective_user.id): return
    try:
        target_id = int(context.args[0])
        ud = context.application.user_data.get(target_id)
        if ud:
            had_sub = "subscription_expires" in ud or "file_subscription_expires" in ud
            ud.pop("subscription_expires", None)
            ud.pop("file_subscription_expires", None)
            await update.message.reply_text("🛑 Все подписки пользователя отменены." if had_sub else "У пользователя нет активных подписок.")
        else:
            await update.message.reply_text("Пользователь не найден.")
    except (IndexError, ValueError):
        await update.message.reply_text("Используйте: /delsub <user_id>")

async def admin_panel_open(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    if not is_admin(update.effective_user.id):
        if update.callback_query: await update.callback_query.answer("Только для администратора.", show_alert=True)
        return MAIN_MENU
    
    text = "⚙️ <b>Админ-панель</b>"
    if update.callback_query:
        await update.callback_query.answer()
        await update.callback_query.message.edit_text(text, parse_mode="HTML", reply_markup=admin_menu_kb())
    else:
        await update.message.reply_html(text, reply_markup=admin_menu_kb())
    return ADMIN_MENU

async def admin_reset_start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await update.callback_query.answer()
    await update.callback_query.message.edit_text("Введите <b>ID</b> для сброса лимитов:", parse_mode="HTML", reply_markup=admin_cancel_kb())
    return ADMIN_RESET_WAIT_ID

async def admin_reset_receive_id(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    try:
        target_id = int(update.message.text)
        ud = context.application.user_data.get(target_id)
        if ud:
            today = _today()
            usage = ud.setdefault("usage", {})
            usage["rewriter"] = {"count": 0, "date": today}
            usage["literature"] = {"count": 0, "date": today}
            usage["file_rewrite"] = {"count": 0, "date": today}
            await update.message.reply_text("✅ Сброшено.", reply_markup=admin_cancel_kb())
        else:
            await update.message.reply_text("Пользователь не найден.", reply_markup=admin_cancel_kb())
    except ValueError:
        await update.message.reply_text("Неверный ID.", reply_markup=admin_cancel_kb())
    return ADMIN_MENU

async def admin_addsub_start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await update.callback_query.answer()
    await update.callback_query.message.edit_text("Введите ID для выдачи общей подписки:", reply_markup=admin_cancel_kb())
    return ADMIN_ADDSUB_WAIT_ID

async def admin_addsub_receive_id(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    try:
        target_id = int(update.message.text)
        context.user_data["admin_target"] = target_id
        await update.message.reply_html(f"ID <code>{target_id}</code>. Введите кол-во дней:", reply_markup=admin_cancel_kb())
        return ADMIN_ADDSUB_WAIT_DAYS
    except ValueError:
        await update.message.reply_text("Неверный ID.", reply_markup=admin_cancel_kb())
        return ADMIN_ADDSUB_WAIT_ID

async def admin_addsub_receive_days(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    target_id = context.user_data.pop("admin_target", None)
    if not target_id:
        await update.message.reply_text("Сессия истекла.", reply_markup=admin_cancel_kb()); return ADMIN_MENU
    try:
        days = int(update.message.text)
        ud = context.application.user_data.get(target_id)
        if not ud:
             await update.message.reply_text("Пользователь не найден.", reply_markup=admin_cancel_kb()); return ADMIN_MENU

        start_date = datetime.now().date()
        exp_str = ud.get("subscription_expires")
        if exp_str:
            try:
                exp_date = datetime.strptime(exp_str, "%Y-%m-%d").date()
                if exp_date > start_date: start_date = exp_date
            except (ValueError, TypeError): pass
        
        new_exp = start_date + timedelta(days=days)
        ud["subscription_expires"] = new_exp.strftime("%Y-%m-%d")
        await update.message.reply_text(f"✅ Подписка выдана до {new_exp.strftime('%d.%m.%Y')}", reply_markup=admin_cancel_kb())
    except (ValueError, KeyError):
        await update.message.reply_text("Ошибка ввода.", reply_markup=admin_cancel_kb())
    return ADMIN_MENU

async def admin_addsub_file_start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await update.callback_query.answer()
    await update.callback_query.message.edit_text("Введите ID для выдачи подписки на файлы:", reply_markup=admin_cancel_kb())
    return ADMIN_ADDSUB_FILE_WAIT_ID

async def admin_addsub_file_receive_id(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    try:
        target_id = int(update.message.text)
        context.user_data["admin_target"] = target_id
        await update.message.reply_html(f"ID <code>{target_id}</code>. Введите кол-во дней:", reply_markup=admin_cancel_kb())
        return ADMIN_ADDSUB_FILE_WAIT_DAYS
    except ValueError:
        await update.message.reply_text("Неверный ID.", reply_markup=admin_cancel_kb())
        return ADMIN_ADDSUB_FILE_WAIT_ID

async def admin_addsub_file_receive_days(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    target_id = context.user_data.pop("admin_target", None)
    if not target_id:
        await update.message.reply_text("Сессия истекла.", reply_markup=admin_cancel_kb()); return ADMIN_MENU
    try:
        days = int(update.message.text)
        ud = context.application.user_data.get(target_id)
        if not ud:
             await update.message.reply_text("Пользователь не найден.", reply_markup=admin_cancel_kb()); return ADMIN_MENU
        
        start_date = datetime.now().date()
        exp_str = ud.get("file_subscription_expires")
        if exp_str:
            try:
                exp_date = datetime.strptime(exp_str, "%Y-%m-%d").date()
                if exp_date > start_date: start_date = exp_date
            except (ValueError, TypeError): pass
            
        new_exp = start_date + timedelta(days=days)
        ud["file_subscription_expires"] = new_exp.strftime("%Y-%m-%d")
        await update.message.reply_text(f"✅ Подписка на файлы выдана до {new_exp.strftime('%d.%m.%Y')}", reply_markup=admin_cancel_kb())
    except (ValueError, KeyError):
        await update.message.reply_text("Ошибка ввода.", reply_markup=admin_cancel_kb())
    return ADMIN_MENU

async def admin_delsub_start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await update.callback_query.answer()
    await update.callback_query.message.edit_text("Введите ID для отмены всех подписок:", reply_markup=admin_cancel_kb())
    return ADMIN_DELSUB_WAIT_ID

async def admin_delsub_receive_id(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    try:
        target_id = int(update.message.text)
        ud = context.application.user_data.get(target_id)
        if ud:
            had_sub = "subscription_expires" in ud or "file_subscription_expires" in ud
            ud.pop("subscription_expires", None)
            ud.pop("file_subscription_expires", None)
            await update.message.reply_text("🛑 Все подписки отменены." if had_sub else "У пользователя нет подписок.", reply_markup=admin_cancel_kb())
        else:
            await update.message.reply_text("Пользователь не найден.", reply_markup=admin_cancel_kb())
    except ValueError:
        await update.message.reply_text("Неверный ID.", reply_markup=admin_cancel_kb())
    return ADMIN_MENU
    
async def admin_status_start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await update.callback_query.answer()
    await update.callback_query.message.edit_text("Введите ID для просмотра статуса:", parse_mode="HTML", reply_markup=admin_cancel_kb())
    return ADMIN_STATUS_WAIT_ID

async def admin_status_receive_id(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    try:
        target_id = int(update.message.text)
        ud = context.application.user_data.get(target_id)
        if not ud:
            await update.message.reply_text("Пользователь не найден.", reply_markup=admin_cancel_kb())
            return ADMIN_MENU
        
        exp_main = ud.get("subscription_expires")
        sub_main = f"до {datetime.strptime(exp_main, '%Y-%m-%d').strftime('%d.%m.%Y')}" if exp_main else "нет"
        
        exp_file = ud.get("file_subscription_expires")
        sub_file = f"до {datetime.strptime(exp_file, '%Y-%m-%d').strftime('%d.%m.%Y')}" if exp_file else "нет"
        
        usage = ud.get("usage", {})
        rew = usage.get("rewriter", {})
        lit = usage.get("literature", {})
        fil = usage.get("file_rewrite", {})
        
        tags = ", ".join(ud.get("tags", [])) or "—"
        refs_set = ud.get("referrals", set())
        refs = len(refs_set) if isinstance(refs_set, set) else 0

        text = (f"👤 <b>{target_id}</b>\n"
                f"Подписка (общая): {sub_main}\n"
                f"Подписка (файлы): {sub_file}\n"
                f"Рерайтер исп: {rew.get('count',0)} ({rew.get('date','-')})\n"
                f"Литература исп: {lit.get('count',0)} ({lit.get('date','-')})\n"
                f"Файлы исп: {fil.get('count',0)} ({fil.get('date','-')})\n"
                f"Теги: {html.escape(tags)}\nРефералов: {refs}")
        await update.message.reply_html(text, reply_markup=admin_cancel_kb())

    except ValueError:
        await update.message.reply_text("Неверный ID.", reply_markup=admin_cancel_kb())
    return ADMIN_MENU

async def admin_search_start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await update.callback_query.answer()
    await update.callback_query.message.edit_text("Введите ID или часть юзернейма:", reply_markup=admin_cancel_kb())
    return ADMIN_SEARCH_WAIT_QUERY

async def admin_search_do(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    q = (update.message.text or "").strip().lower()
    res = []
    if q.isdigit():
        uid = int(q)
        if uid in context.application.user_data:
            u = context.application.user_data[uid]
            res.append((uid, u.get("last_username",""), u.get("first_seen","?")))
    else:
        for uid, data in context.application.user_data.items():
            if q in (data.get("last_username","") or "").lower():
                res.append((uid, data.get("last_username",""), data.get("first_seen","?")))
            if len(res) >= 20: break
    if not res:
        await update.message.reply_text("Ничего не найдено.", reply_markup=admin_cancel_kb())
    else:
        lines = [f"• <code>{uid}</code> @{name or '—'} (с {fs})" for uid,name,fs in res]
        await update.message.reply_html("Найдено:\n" + "\n".join(lines), reply_markup=admin_cancel_kb())
    return ADMIN_MENU
    
async def admin_tags_start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await update.callback_query.answer()
    await update.callback_query.message.edit_text("Введите ID пользователя:", reply_markup=admin_cancel_kb())
    return ADMIN_TAGS_WAIT_ID

async def admin_tags_id(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    try:
        target = int(update.message.text)
        context.user_data["admin_target"] = target
        await update.message.reply_text("Введите теги через запятую (-тег для удаления):\nпример: vip, -test", reply_markup=admin_cancel_kb())
        return ADMIN_TAGS_WAIT_VALUE
    except ValueError:
        await update.message.reply_text("Неверный ID.", reply_markup=admin_cancel_kb())
        return ADMIN_TAGS_WAIT_ID

async def admin_tags_set(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    target = context.user_data.pop("admin_target", None)
    if not target: await update.message.reply_text("Сессия истекла.", reply_markup=admin_cancel_kb()); return ADMIN_MENU
    
    ud = context.application.user_data.get(target)
    if not ud: await update.message.reply_text("Пользователь не найден.", reply_markup=admin_cancel_kb()); return ADMIN_MENU
    
    tags_input = [t.strip() for t in (update.message.text or "").split(",") if t.strip()]
    cur = set(ud.get("tags", []))
    for t in tags_input:
        if t.startswith("-"): cur.discard(t[1:].strip())
        else: cur.add(t)
    ud["tags"] = sorted(list(cur))
    await update.message.reply_text(f"Готово. Теги: {', '.join(ud['tags']) or '—'}", reply_markup=admin_cancel_kb())
    return ADMIN_MENU

async def admin_broadcast_start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await update.callback_query.answer()
    await update.callback_query.message.edit_text("Выберите сегмент рассылки:", reply_markup=broadcast_segments_kb())
    return ADMIN_BROADCAST_SEGMENT

async def admin_broadcast_pick(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await update.callback_query.answer()
    data = update.callback_query.data
    context.user_data["b_silent"] = data.endswith("_silent")
    context.user_data["b_segment"] = data.replace("_silent", "")
    mode = " (тихая)" if context.user_data["b_silent"] else ""
    await update.callback_query.message.edit_text(f"Введите текст рассылки{mode}:", reply_markup=admin_cancel_kb())
    return ADMIN_BROADCAST_WAIT_TEXT

async def admin_broadcast_send(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    seg = context.user_data.pop("b_segment", None)
    silent = context.user_data.pop("b_silent", False)
    txt = update.message.text
    if not seg or not txt:
        await update.message.reply_text("Ошибка.", reply_markup=admin_cancel_kb()); return ADMIN_MENU

    uids = set()
    cutoff = datetime.now() - timedelta(days=7)
    for uid, data in context.application.user_data.items():
        is_sub = False
        exp_str = data.get("subscription_expires")
        if exp_str:
            try: is_sub = datetime.strptime(exp_str, "%Y-%m-%d").date() >= datetime.now().date()
            except Exception: pass

        if seg == "bseg_all": uids.add(uid)
        elif seg == "bseg_free" and not is_sub: uids.add(uid)
        elif seg == "bseg_subs" and is_sub: uids.add(uid)
        elif seg == "bseg_inactive":
            ls_str = data.get("last_seen")
            if ls_str:
                try:
                    if datetime.strptime(ls_str, "%Y-%m-%d %H:%M:%S") < cutoff: uids.add(uid)
                except Exception: pass
    
    sent = 0
    await update.message.reply_text(f"Начинаю рассылку для {len(uids)} пользователей...")
    for uid in uids:
        try:
            await context.bot.send_message(chat_id=uid, text=txt, disable_notification=silent)
            sent += 1
            await asyncio.sleep(0.1)
        except Exception:
            continue
    await update.message.reply_text(f"Готово. Отправлено: {sent}", reply_markup=admin_cancel_kb())
    return ADMIN_MENU

async def admin_export_csv(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await update.callback_query.answer()
    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow(["id", "first_seen", "last_seen", "username", "sub_expires", "file_sub_expires", "tags", "refs"])
    for uid, d in context.application.user_data.items():
        refs = d.get("referrals", set())
        w.writerow([uid, d.get("first_seen",""), d.get("last_seen",""), d.get("last_username",""),
                    d.get("subscription_expires",""), d.get("file_subscription_expires",""),
                    "|".join(d.get("tags", [])), len(refs) if isinstance(refs, set) else 0])
    byte = io.BytesIO(buf.getvalue().encode("utf-8")); byte.name = "users.csv"
    await update.callback_query.message.reply_document(InputFile(byte))
    return ADMIN_MENU

async def admin_blacklist_start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await update.callback_query.answer()
    await update.callback_query.message.edit_text("Введите ID для ЧС (-ID для удаления):", reply_markup=admin_cancel_kb())
    return ADMIN_BLACKLIST_WAIT_ID

async def admin_blacklist_apply(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    txt = update.message.text.strip()
    bl = context.application.bot_data.setdefault("blacklist", set())
    try:
        if txt.startswith("-"):
            bl.discard(int(txt[1:]))
            await update.message.reply_text("Удалён из ЧС.", reply_markup=admin_cancel_kb())
        else:
            bl.add(int(txt))
            await update.message.reply_text("Добавлен в ЧС.", reply_markup=admin_cancel_kb())
    except ValueError:
        await update.message.reply_text("Неверный формат.", reply_markup=admin_cancel_kb())
    return ADMIN_MENU

async def admin_shadow_start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await update.callback_query.answer()
    await update.callback_query.message.edit_text("Введите ID для теневого бана (-ID для снятия):", reply_markup=admin_cancel_kb())
    return ADMIN_SHADOW_WAIT_ID

async def admin_shadow_apply(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    txt = update.message.text.strip()
    sb = context.application.bot_data.setdefault("shadowban", set())
    try:
        if txt.startswith("-"):
            sb.discard(int(txt[1:]))
            await update.message.reply_text("Теневой бан снят.", reply_markup=admin_cancel_kb())
        else:
            sb.add(int(txt))
            await update.message.reply_text("Теневой бан установлен.", reply_markup=admin_cancel_kb())
    except ValueError:
        await update.message.reply_text("Неверный формат.", reply_markup=admin_cancel_kb())
    return ADMIN_MENU

async def admin_setlimit_start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await update.callback_query.answer()
    await update.callback_query.message.edit_text("Введите ID пользователя:", reply_markup=admin_cancel_kb())
    return ADMIN_SETLIMIT_WAIT_ID

async def admin_setlimit_id(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    try:
        target = int(update.message.text)
        context.user_data["admin_target"] = target
        await update.message.reply_text("Введите лимиты через пробел (рерайтер литература файлы):\nнапр. 2 1 0", reply_markup=admin_cancel_kb())
        return ADMIN_SETLIMIT_WAIT_VALUES
    except ValueError:
        await update.message.reply_text("Неверный ID.", reply_markup=admin_cancel_kb())
        return ADMIN_SETLIMIT_WAIT_ID

async def admin_setlimit_values(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    target = context.user_data.pop("admin_target", None)
    if not target: await update.message.reply_text("Сессия истекла.", reply_markup=admin_cancel_kb()); return ADMIN_MENU
    try:
        parts = update.message.text.split()
        if len(parts) != 3: raise ValueError("Must be 3 values")
        c1, c2, c3 = map(int, parts)
        ud = context.application.user_data.get(target)
        if not ud: await update.message.reply_text("Пользователь не найден.", reply_markup=admin_cancel_kb()); return ADMIN_MENU
        
        today = _today()
        u = ud.setdefault("usage", {})
        u["rewriter"] = {"count": max(0, c1), "date": today}
        u["literature"] = {"count": max(0, c2), "date": today}
        u["file_rewrite"] = {"count": max(0, c3), "date": today}
        await update.message.reply_text("✅ Установлено.", reply_markup=admin_cancel_kb())
    except ValueError:
        await update.message.reply_text("Нужно 3 числа через пробел.", reply_markup=admin_cancel_kb())
        return ADMIN_SETLIMIT_WAIT_VALUES
    return ADMIN_MENU

async def admin_metrics_open(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await update.callback_query.answer()
    bd = context.application.bot_data
    text = (
        f"📈 <b>Метрики</b> ({_today()} {_now_hms()})\n\n"
        f"• Всего уникальных: <b>{len(bd.get('all_users', set()))}</b>\n"
        f"• DAU: <b>{_count_dau(bd, 1)}</b>, WAU: <b>{_count_dau(bd, 7)}</b>, MAU: <b>{_count_dau(bd, 30)}</b>\n"
        f"• Новых сегодня: <b>{_count_new_users(context.application, 1)}</b>, за 7д: <b>{_count_new_users(context.application, 7)}</b>\n\n"
        f"• Использование сегодня:\n"
        f"  ├─ Рерайтер: <b>{_feature_usage_today(bd).get('rewriter',0)}</b>\n"
        f"  ├─ Литература: <b>{_feature_usage_today(bd).get('literature',0)}</b>\n"
        f"  └─ Рерайт файлов: <b>{_feature_usage_today(bd).get('file_rewrite',0)}</b>"
    )
    await update.callback_query.message.edit_text(text, parse_mode="HTML", reply_markup=admin_metrics_kb())
    return ADMIN_METRICS_MENU

async def admin_metrics_refresh(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    return await admin_metrics_open(update, context)

async def admin_metrics_export(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await update.callback_query.answer()
    bd = context.application.bot_data
    dau, feat = bd.get("dau", {}), bd.get("feature_usage_today", {})
    buf = io.StringIO(); w = csv.writer(buf)
    w.writerow(["date", "dau", "rewriter", "literature", "file_rewrite", "new_users"])
    for i in range(30):
        d = datetime.now().date() - timedelta(days=i)
        key = d.strftime("%Y-%m-%d")
        f = feat.get(key, {})
        new_cnt = 0
        for _, ud in context.application.user_data.items():
            if ud.get("first_seen") == key: new_cnt += 1
        w.writerow([key, len(dau.get(key, set())), f.get("rewriter", 0), f.get("literature", 0), f.get("file_rewrite", 0), new_cnt])
    byte = io.BytesIO(buf.getvalue().encode("utf-8")); byte.name = "metrics.csv"
    await update.callback_query.message.reply_document(InputFile(byte))
    return ADMIN_METRICS_MENU

async def admin_maint_toggle(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await update.callback_query.answer()
    cur = _maintenance_on(context.application)
    context.application.bot_data["maintenance_enabled"] = not cur
    status = "ВКЛЮЧЕН" if not cur else "ВЫКЛЮЧЕН"
    text = (f"🛠 Режим техработ: <b>{status}</b>.\n\n"
            f"Текст для пользователей:\n{_maintenance_text(context.application)}")
    await update.callback_query.message.edit_text(
        text, parse_mode="HTML", reply_markup=admin_menu_kb(), disable_web_page_preview=True
    )
    return ADMIN_MENU
    
async def admin_maint_msg_start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await update.callback_query.answer()
    await update.callback_query.message.edit_text("Введите новый текст для тех. работ (HTML):", reply_markup=admin_cancel_kb())
    return ADMIN_MAINT_MSG_WAIT

async def admin_maint_msg_save(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    txt = update.message.text
    if txt:
        context.application.bot_data["maintenance_msg"] = txt[:1000]
        await update.message.reply_html("✅ Текст сохранён.", reply_markup=admin_cancel_kb())
    else:
        await update.message.reply_text("Текст не изменён.", reply_markup=admin_cancel_kb())
    return ADMIN_MENU

async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await update.effective_message.reply_text("Действие отменено.")
    return await start(update, context)

async def error_handler(update: object, context: ContextTypes.DEFAULT_TYPE) -> None:
    logger.error("Exception while handling an update:", exc_info=context.error)
    if isinstance(update, Update) and update.effective_message:
        try:
            await update.effective_message.reply_text("Произошла ошибка. Попробуйте /start")
        except Exception as e:
            logger.error(f"Failed to send error message: {e}")

# ===== MAIN BOT SETUP =====
def main() -> None:
    if not BOT_TOKEN:
        raise ValueError("BOT_TOKEN не задан в переменных окружения.")

    persistence = PicklePersistence(filepath="bot_persistence")
    app = Application.builder().token(BOT_TOKEN).persistence(persistence).build()

    app.add_error_handler(error_handler)
    app.add_handler(CommandHandler("help", help_cmd))
    app.add_handler(CommandHandler("service", service_cmd))
    app.add_handler(CommandHandler("report", report_cmd))
    app.add_handler(CommandHandler("status", check_status))
    app.add_handler(CommandHandler("reset", reset_limit))
    app.add_handler(CommandHandler("addsub", add_subscription))
    app.add_handler(CommandHandler("delsub", del_subscription))
    app.add_handler(CommandHandler("metrics", admin_metrics_open))


    conv_handler = ConversationHandler(
        entry_points=[CommandHandler("start", start)],
        states={
            MAIN_MENU: [
                CallbackQueryHandler(rewriter_start, pattern="^rewriter$"),
                CallbackQueryHandler(file_rewriter_start, pattern="^file_rewriter$"),
                CallbackQueryHandler(literature_start, pattern="^literature$"),
                CallbackQueryHandler(gost_menu, pattern="^gost$"),
                CallbackQueryHandler(cabinet_open, pattern="^cabinet$"),
                CallbackQueryHandler(admin_panel_open, pattern="^admin_panel$"),
                CallbackQueryHandler(start, pattern="^back_to_main_menu$"),
            ],
            REWRITER_TEXT_INPUT: [MessageHandler(filters.TEXT & ~filters.COMMAND, rewriter_process_text)],
            LITERATURE_TOPIC_INPUT: [MessageHandler(filters.TEXT & ~filters.COMMAND, literature_process_topic)],
            FILE_REWRITE_WAIT_FILE: [MessageHandler(filters.Document.ALL, process_document_rewrite)],
            GOST_MENU: [
                CallbackQueryHandler(gost_show_info, pattern="^gost_(footnote|references|general)$"),
                CallbackQueryHandler(gost_menu, pattern="^gost_back$"),
            ],
            CABINET_MENU: [
                CallbackQueryHandler(cabinet_history, pattern="^cab_history$"),
                CallbackQueryHandler(cabinet_export, pattern="^cab_export$"),
                CallbackQueryHandler(cabinet_refresh, pattern="^cab_refresh$"),
                CallbackQueryHandler(settings_open, pattern="^cab_settings$"),
                CallbackQueryHandler(cabinet_referrals, pattern="^cab_ref$"),
            ],
            CAPTCHA_WAIT: [MessageHandler(filters.TEXT & ~filters.COMMAND, captcha_check)],
            SETTINGS_MENU: [CallbackQueryHandler(settings_pick, pattern="^(set_tone|set_gost)$")],
            SETTINGS_TONE_WAIT: [MessageHandler(filters.TEXT & ~filters.COMMAND, settings_tone_save)],
            SETTINGS_GOST_WAIT: [MessageHandler(filters.TEXT & ~filters.COMMAND, settings_gost_save)],
            ADMIN_MENU: [
                CallbackQueryHandler(admin_reset_start, pattern="^admin_reset$"),
                CallbackQueryHandler(admin_addsub_start, pattern="^admin_addsub$"),
                CallbackQueryHandler(admin_addsub_file_start, pattern="^admin_addsub_file$"),
                CallbackQueryHandler(admin_delsub_start, pattern="^admin_delsub$"),
                CallbackQueryHandler(admin_status_start, pattern="^admin_status$"),
                CallbackQueryHandler(admin_search_start, pattern="^admin_search$"),
                CallbackQueryHandler(admin_tags_start, pattern="^admin_tags$"),
                CallbackQueryHandler(admin_broadcast_start, pattern="^admin_broadcast$"),
                CallbackQueryHandler(admin_export_csv, pattern="^admin_export$"),
                CallbackQueryHandler(admin_blacklist_start, pattern="^admin_blacklist$"),
                CallbackQueryHandler(admin_shadow_start, pattern="^admin_shadow$"),
                CallbackQueryHandler(admin_setlimit_start, pattern="^admin_setlimit$"),
                CallbackQueryHandler(admin_metrics_open, pattern="^admin_metrics$"),
                CallbackQueryHandler(admin_maint_toggle, pattern="^admin_maint_toggle$"),
                CallbackQueryHandler(admin_maint_msg_start, pattern="^admin_maint_msg$"),
                CallbackQueryHandler(start, pattern="^back_to_main_menu$"),
            ],
            ADMIN_RESET_WAIT_ID: [MessageHandler(filters.TEXT & ~filters.COMMAND, admin_reset_receive_id)],
            ADMIN_ADDSUB_WAIT_ID: [MessageHandler(filters.TEXT & ~filters.COMMAND, admin_addsub_receive_id)],
            ADMIN_ADDSUB_WAIT_DAYS: [MessageHandler(filters.TEXT & ~filters.COMMAND, admin_addsub_receive_days)],
            ADMIN_STATUS_WAIT_ID: [MessageHandler(filters.TEXT & ~filters.COMMAND, admin_status_receive_id)],
            ADMIN_DELSUB_WAIT_ID: [MessageHandler(filters.TEXT & ~filters.COMMAND, admin_delsub_receive_id)],
            ADMIN_SEARCH_WAIT_QUERY: [MessageHandler(filters.TEXT & ~filters.COMMAND, admin_search_do)],
            ADMIN_TAGS_WAIT_ID: [MessageHandler(filters.TEXT & ~filters.COMMAND, admin_tags_id)],
            ADMIN_TAGS_WAIT_VALUE: [MessageHandler(filters.TEXT & ~filters.COMMAND, admin_tags_set)],
            ADMIN_BROADCAST_SEGMENT: [CallbackQueryHandler(admin_broadcast_pick, pattern="^bseg_")],
            ADMIN_BROADCAST_WAIT_TEXT: [MessageHandler(filters.TEXT & ~filters.COMMAND, admin_broadcast_send)],
            ADMIN_SETLIMIT_WAIT_ID: [MessageHandler(filters.TEXT & ~filters.COMMAND, admin_setlimit_id)],
            ADMIN_SETLIMIT_WAIT_VALUES: [MessageHandler(filters.TEXT & ~filters.COMMAND, admin_setlimit_values)],
            ADMIN_BLACKLIST_WAIT_ID: [MessageHandler(filters.TEXT & ~filters.COMMAND, admin_blacklist_apply)],
            ADMIN_SHADOW_WAIT_ID: [MessageHandler(filters.TEXT & ~filters.COMMAND, admin_shadow_apply)],
            ADMIN_METRICS_MENU: [
                CallbackQueryHandler(admin_metrics_refresh, pattern="^admin_metrics_refresh$"),
                CallbackQueryHandler(admin_metrics_export, pattern="^admin_metrics_export$"),
                CallbackQueryHandler(admin_panel_open, pattern="^admin_panel$"),
            ],
            ADMIN_ADDSUB_FILE_WAIT_ID: [MessageHandler(filters.TEXT & ~filters.COMMAND, admin_addsub_file_receive_id)],
            ADMIN_ADDSUB_FILE_WAIT_DAYS: [MessageHandler(filters.TEXT & ~filters.COMMAND, admin_addsub_file_receive_days)],
            ADMIN_MAINT_MSG_WAIT: [MessageHandler(filters.TEXT & ~filters.COMMAND, admin_maint_msg_save)],
        },
        fallbacks=[CommandHandler("cancel", cancel), CommandHandler("start", start)],
        allow_reentry=True,
        persistent=True,
        name="main_conversation"
    )
    app.add_handler(conv_handler)
    
    logger.info("Бот запущен...")
    app.run_polling()


if __name__ == "__main__":
    main()

