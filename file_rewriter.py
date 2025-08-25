# -*- coding: utf-8 -*-

from __future__ import annotations

from dataclasses import dataclass
from typing import List, Tuple, Callable, Awaitable, Optional, Any, TYPE_CHECKING
import io
from collections import Counter

# ------------------------- docx imports (без ругани Pylance) -------------------------
try:
    import docx  # python-docx
    from docx.enum.text import WD_COLOR_INDEX
except ImportError:  # библиотека не установлена — дадим понятную ошибку позднее
    docx = None  # type: ignore
    WD_COLOR_INDEX = None  # type: ignore

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument
    from docx.text.paragraph import Paragraph
    from docx.text.run import Run
else:
    DocxDocument = Any  # type: ignore
    Paragraph = Any  # type: ignore
    Run = Any  # type: ignore


# ================================== DATA ==================================

@dataclass
class HighlightedPart:
    """Фрагмент, выделенный заливкой (желтой), собранный из нескольких runs."""
    original_text: str
    part_index: int
    run_indices: List[Tuple[int, int]]  # список (paragraph_index, run_index)


@dataclass
class RewrittenPart:
    """Результат рерайта определенного фрагмента."""
    rewritten_text: str
    original_part: HighlightedPart


# ============================ PARSE HIGHLIGHT ==============================

def process_docx_for_rewrite(file_bytes: bytes) -> Tuple[DocxDocument, List[HighlightedPart]]:
    """
    Открывает .docx и находит фрагменты, выделенные заливкой YELLOW.
    Возвращает сам документ и список выделенных частей (в порядке следования).
    """
    if docx is None:
        raise RuntimeError("Библиотека python-docx не установлена. Выполните: pip install python-docx")

    document: DocxDocument = docx.Document(io.BytesIO(file_bytes))  # type: ignore[call-arg]
    highlighted: List[HighlightedPart] = []

    current_text: List[str] = []
    current_runs: List[Tuple[int, int]] = []
    part_idx = 0

    for p_i, p in enumerate(document.paragraphs):
        for r_i, run in enumerate(p.runs):
            is_hl = (getattr(run.font, "highlight_color", None) == getattr(WD_COLOR_INDEX, "YELLOW", None))
            if is_hl:
                current_text.append(run.text)
                current_runs.append((p_i, r_i))
            else:
                if current_runs:
                    highlighted.append(
                        HighlightedPart(
                            original_text="".join(current_text),
                            part_index=part_idx,
                            run_indices=current_runs[:],
                        )
                    )
                    part_idx += 1
                    current_text.clear()
                    current_runs.clear()

    # хвост
    if current_runs:
        highlighted.append(
            HighlightedPart(
                original_text="".join(current_text),
                part_index=part_idx,
                run_indices=current_runs[:],
            )
        )

    return document, highlighted


# ============================ PROMPTS / UTILS ==============================

def _chunk_text(text: str, max_chars: int = 8000) -> List[str]:
    s = (text or "").strip()
    if len(s) <= max_chars:
        return [s]
    return [s[i:i + max_chars] for i in range(0, len(s), max_chars)]


def _build_prompt(
    chunk: str,
    tone: str,
    target_uniqueness: str,
    deep_analyze: bool = True
) -> str:
    """
    Инструкции «внутри промпта»:
    - сначала (мысленно) анализ смысла, затем рерайт;
    - только перефразированный текст на выходе, без пояснений/списков/предисловий.
    - под антиплагиат: синтаксические перестройки, синонимизация, изменение порядка, но без «воды».
    """
    uniq_line = f"Цель по проверке оригинальности (антиплагиат): {target_uniqueness}.\n" if target_uniqueness else ""
    analyze = (
        "Сначала мысленно выдели ключевые идеи, термины и логические связи, чтобы понять смысл. "
        "Не показывай ход анализа пользователю. Затем перефразируй.\n"
        if deep_analyze else ""
    )

    return (
        "Ты — академический редактор на русском. Перепиши фрагмент строго по правилам:\n"
        f"{analyze}"
        f"- стиль: {tone};\n"
        "- сохрани исходную мысль, структуру абзацев и порядок смысловых блоков;\n"
        "- не добавляй списки/заголовки/вводные, если их нет в тексте;\n"
        "- не вставляй выводы, пояснения, комментарии, метатекст;\n"
        "- перефразируй так, чтобы снизить совпадения: меняй синтаксис, перестраивай фразы, используй синонимы, "
        "инвертируй порядок, дроби/сливай предложения — но без потери смысла;\n"
        "- избегай шаблонов ИИ, штампов и речевых клише; соблюдай академическую норму русского языка;\n"
        f"{uniq_line}"
        "ВЫВЕДИ ТОЛЬКО перефразированный текст, без кавычек, без пояснений.\n\n"
        "ТЕКСТ:\n"
        f"{chunk}"
    )


def _is_ai_error_text(s: str) -> bool:
    """Эвристика: похоже ли это на сообщение об ошибке/квоте/перегрузке?"""
    s = (s or "").strip()
    if not s:
        return True
    low = s.lower()
    return any(kw in low for kw in (
        "ошибка", "не удалось", "недоступен", "перегруз", "quota", "rate limit",
        "429", "api key", "try again", "попробуйте позже"
    ))


# =========================== REWRITE PIPELINE ==============================

async def rewrite_highlighted_parts_async(
    parts: List[HighlightedPart],
    rewrite_fn: Callable[[str], Awaitable[str]],
    tone: str,
    target_uniqueness: str,
    deep_analyze: bool = True,
) -> List[RewrittenPart]:
    """
    Для каждой выделенной части вызывает ИИ. Если ИИ вернул ошибку — оставляем исходный текст.
    """
    out: List[RewrittenPart] = []

    for part in parts:
        src = part.original_text or ""
        if not src.strip():
            continue

        chunks = _chunk_text(src)
        rewritten_chunks: List[str] = []

        for ch in chunks:
            prompt = _build_prompt(ch, tone=tone, target_uniqueness=target_uniqueness, deep_analyze=deep_analyze)
            resp = (await rewrite_fn(prompt)).strip()

            if _is_ai_error_text(resp):
                # ИИ не смог — возвращаем исходный кусок без изменений
                rewritten_chunks.append(ch)
            else:
                # убираем возможные бэктики/служебные символы
                rewritten_chunks.append(resp.strip("`").strip())

        out.append(
            RewrittenPart(
                rewritten_text=" ".join(rewritten_chunks).strip(),
                original_part=part
            )
        )

    return out


# ============================== APPLY BACK ================================

def _majority_style_attrs(original_doc: DocxDocument, run_refs: List[Tuple[int, int]]) -> dict:
    """
    Берём преобладающие атрибуты шрифта на выделенном участке — чтобы не сделать весь абзац жирным,
    если только первый run был жирным.
    """
    bold_vals: List[Optional[bool]] = []
    italic_vals: List[Optional[bool]] = []
    underline_vals: List[Optional[bool]] = []
    sizes: List[Any] = []  # длина (EMU) из python-docx, тип оставляем Any

    for p_i, r_i in run_refs:
        run: Run = original_doc.paragraphs[p_i].runs[r_i]
        bold_vals.append(run.bold)
        italic_vals.append(run.italic)
        underline_vals.append(run.underline)
        if run.font and run.font.size:
            sizes.append(run.font.size)

    def _maj(lst: List[Optional[bool]]) -> Optional[bool]:
        # считаем только True/False; None трактуем как отсутствие явной установки
        vals = [v for v in lst if isinstance(v, bool)]
        if not vals:
            return None
        cnt = Counter(vals)
        # если большинство False — возвращаем False, иначе True
        return cnt.most_common(1)[0][0]

    def _mode(lst: List[Any]) -> Optional[Any]:
        return Counter(lst).most_common(1)[0][0] if lst else None

    return {
        "bold": _maj(bold_vals),
        "italic": _maj(italic_vals),
        "underline": _maj(underline_vals),
        "size": _mode(sizes),
    }


def build_final_docx(original_doc: DocxDocument, rewritten_parts: List[RewrittenPart]) -> bytes:
    """
    Вставляет переписанные фрагменты:
    - если текст не изменился (или была ошибка ИИ) — ничего не трогаем;
    - если изменился — подставляем в первый run участка новый текст, сбрасываем подсветку
      и ставим преобладающие атрибуты шрифта (bold/italic/underline/size), чтобы не ломать оформление.
    Остальные runs в участке очищаем.
    """
    parts_map = {p.original_part.part_index: p for p in rewritten_parts}

    for part in parts_map.values():
        orig = (part.original_part.original_text or "").strip()
        new = (part.rewritten_text or "").strip()

        # Ничего не меняем, если реального рерайта нет
        if not new or new == orig:
            continue

        # Куда вставлять
        if not part.original_part.run_indices:
            continue
        p0, r0 = part.original_part.run_indices[0]
        first_run: Run = original_doc.paragraphs[p0].runs[r0]

        # выставим преобладающие атрибуты стиля на участке
        style = _majority_style_attrs(original_doc, part.original_part.run_indices)

        # вставка текста + аккуратный стиль
        first_run.text = new
        if getattr(first_run, "font", None) is not None:
            try:
                first_run.font.highlight_color = None
            except Exception:
                pass

        # атрибуты run уровня (bold/italic/underline)
        if style["bold"] is not None:
            first_run.bold = style["bold"]
        if style["italic"] is not None:
            first_run.italic = style["italic"]
        if style["underline"] is not None:
            first_run.underline = style["underline"]
        if style["size"] is not None and getattr(first_run, "font", None) is not None:
            first_run.font.size = style["size"]

        # очищаем остальные runs участка
        for i, (p_i, r_i) in enumerate(part.original_part.run_indices):
            if i == 0:
                continue
            run_to_clear: Run = original_doc.paragraphs[p_i].runs[r_i]
            run_to_clear.text = ""

    # финальный байтовый docx
    bio = io.BytesIO()
    original_doc.save(bio)
    bio.seek(0)
    return bio.getvalue()
